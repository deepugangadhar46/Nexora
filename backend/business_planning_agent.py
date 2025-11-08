"""
BUSINESS PLANNING AGENT for Nexora AI IDE
==========================================

A professional AI agent that turns ideas into structured, fundable business plans using:
- Groq API (Llama model) for all AI features
- Free APIs for regulatory compliance and document export
- Auto Lean Canvas Generator
- Financial Projection Estimator
- Team Role Mapping
- Marketing Strategy Builder
- Investor Summary Generator
- Regulatory Compliance Checker
- Export to DOCX/PDF/Notion

Author: NEXORA Team
Version: 1.0.0
License: MIT
"""

import os
import json
import uuid
import asyncio
import logging
import re
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict, field
from io import BytesIO
import base64

# Third-party imports
import aiohttp
import requests
from dotenv import load_dotenv

# PDF Generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. PDF generation will be disabled.")

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX generation will be disabled.")

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class LeanCanvasBlock:
    """Lean Canvas block data"""
    title: str
    content: List[str]
    description: str = ""


@dataclass
class LeanCanvas:
    """Complete Lean Canvas model"""
    problem: LeanCanvasBlock
    solution: LeanCanvasBlock
    unique_value_proposition: LeanCanvasBlock
    unfair_advantage: LeanCanvasBlock
    customer_segments: LeanCanvasBlock
    key_metrics: LeanCanvasBlock
    channels: LeanCanvasBlock
    cost_structure: LeanCanvasBlock
    revenue_streams: LeanCanvasBlock


@dataclass
class FinancialProjection:
    """Financial projection data"""
    year: int
    revenue: float
    costs: float
    profit: float
    burn_rate: float
    runway_months: int


@dataclass
class FinancialEstimate:
    """3-year financial estimates"""
    projections: List[FinancialProjection]
    cac: float  # Customer Acquisition Cost
    ltv: float  # Lifetime Value
    ltv_cac_ratio: float
    break_even_month: int
    total_funding_needed: float
    assumptions: List[str]


@dataclass
class TeamRole:
    """Team role specification"""
    role: str
    responsibilities: List[str]
    skills_required: List[str]
    priority: str  # Critical/High/Medium/Low
    hiring_timeline: str


@dataclass
class TeamComposition:
    """Complete team structure"""
    roles: List[TeamRole]
    total_team_size: int
    estimated_payroll_monthly: float


@dataclass
class MarketingChannel:
    """Marketing channel details"""
    channel: str
    strategy: str
    estimated_cost: float
    expected_roi: float
    timeline: str


@dataclass
class MarketingStrategy:
    """Complete marketing strategy"""
    channels: List[MarketingChannel]
    total_budget: float
    customer_acquisition_strategy: str
    retention_strategy: str
    growth_tactics: List[str]


@dataclass
class InvestorSummary:
    """One-page investor summary"""
    elevator_pitch: str
    problem_statement: str
    solution_overview: str
    market_opportunity: str
    business_model: str
    traction: str
    competitive_advantage: str
    team_highlights: str
    financial_ask: str
    use_of_funds: List[str]
    key_milestones: List[str]


@dataclass
class ComplianceRequirement:
    """Regulatory compliance requirement"""
    category: str
    requirement: str
    jurisdiction: str
    priority: str  # Critical/High/Medium/Low
    estimated_cost: float
    timeline: str
    resources: List[str]


@dataclass
class RegulatoryCompliance:
    """Complete regulatory compliance analysis"""
    requirements: List[ComplianceRequirement]
    total_compliance_cost: float
    critical_deadlines: List[str]
    recommended_actions: List[str]


@dataclass
class CoFounderFeedback:
    """AI Co-Founder interactive feedback"""
    feedback_type: str  # suggestion/warning/question
    message: str
    reasoning: str
    action_items: List[str]


@dataclass
class BusinessPlanResponse:
    """Complete business plan response"""
    plan_id: str
    business_name: str
    tagline: str
    executive_summary: str
    lean_canvas: LeanCanvas
    financial_estimate: FinancialEstimate
    team_composition: TeamComposition
    marketing_strategy: MarketingStrategy
    investor_summary: InvestorSummary
    regulatory_compliance: RegulatoryCompliance
    co_founder_feedback: List[CoFounderFeedback]
    pdf_url: Optional[str] = None
    docx_url: Optional[str] = None
    notion_url: Optional[str] = None
    created_at: str = ""


# ============================================================================
# API CLIENTS
# ============================================================================

class GroqClient:
    """Groq API client for Llama model"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key or os.getenv("GROQ_API_KEY")
        self.base_url = "https://api.groq.com/openai/v1/chat/completions"
        self.model = "llama-3.3-70b-versatile"  # Using Llama 3.3 70B
        
        if not self.api_key:
            error_msg = (
                "GROQ_API_KEY not found in environment variables. "
                "Please add GROQ_API_KEY=your_key to backend/.env file. "
                "Get your key from https://console.groq.com/"
            )
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        logger.info(f"GroqClient initialized with model: {self.model}")
    
    async def generate(
        self,
        prompt: str,
        system_prompt: Optional[str] = None,
        temperature: float = 0.7,
        max_tokens: int = 4000,
        json_mode: bool = False
    ) -> str:
        """Generate response using Groq Llama model"""
        
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        if json_mode:
            payload["response_format"] = {"type": "json_object"}
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    self.base_url,
                    json=payload,
                    headers=headers,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status != 200:
                        error_text = await response.text()
                        logger.error(f"Groq API error: {error_text}")
                        raise Exception(f"Groq API error: {response.status}")
                    
                    data = await response.json()
                    return data["choices"][0]["message"]["content"]
        
        except Exception as e:
            logger.error(f"Error calling Groq API: {str(e)}")
            raise


class RegulatoryComplianceAPI:
    """Free API for regulatory compliance checking"""
    
    def __init__(self):
        # Using a mock/free regulatory database
        # In production, you could use APIs like:
        # - OpenFDA API (for FDA regulations)
        # - SEC EDGAR API (for financial regulations)
        # - EU Open Data Portal (for EU regulations)
        self.base_url = "https://api.regulations.gov/v4"
    
    async def check_compliance(
        self,
        business_type: str,
        industry: str,
        region: str
    ) -> List[Dict[str, Any]]:
        """Check regulatory requirements for a business"""
        
        # For now, we'll use Groq to generate compliance requirements
        # In production, integrate with actual regulatory APIs
        logger.info(f"Checking compliance for {business_type} in {industry} ({region})")
        
        # Return mock data structure
        return []


class NotionClient:
    """Notion API client for exporting business plans"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key or os.getenv("NOTION_API_KEY")
        self.base_url = "https://api.notion.com/v1"
        self.version = "2022-06-28"
    
    async def create_page(
        self,
        parent_page_id: str,
        title: str,
        content: Dict[str, Any]
    ) -> Optional[str]:
        """Create a Notion page with business plan content"""
        
        if not self.api_key:
            logger.warning("NOTION_API_KEY not set, skipping Notion export")
            return None
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "Notion-Version": self.version
        }
        
        # Build Notion page structure
        payload = {
            "parent": {"page_id": parent_page_id},
            "properties": {
                "title": {
                    "title": [{"text": {"content": title}}]
                }
            },
            "children": []
        }
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.base_url}/pages",
                    json=payload,
                    headers=headers,
                    timeout=aiohttp.ClientTimeout(total=30)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        return data.get("url", "")
                    else:
                        error_text = await response.text()
                        logger.error(f"Notion API error: {error_text}")
                        return None
        
        except Exception as e:
            logger.error(f"Error creating Notion page: {str(e)}")
            return None


# ============================================================================
# BUSINESS PLANNING AGENT
# ============================================================================

class BusinessPlanningAgent:
    """
    BUSINESS PLANNING AGENT
    
    Turns ideas into structured, fundable business plans with:
    - Auto Lean Canvas Generator
    - Financial Projection Estimator
    - Team Role Mapping
    - Marketing Strategy Builder
    - Investor Summary Generator
    - Regulatory Compliance Checker
    - Export to DOCX/PDF/Notion
    - AI Co-Founder Chat
    """
    
    def __init__(
        self,
        groq_api_key: Optional[str] = None,
        notion_api_key: Optional[str] = None
    ):
        """Initialize the Business Planning Agent"""
        
        self.groq = GroqClient(groq_api_key)
        self.compliance_api = RegulatoryComplianceAPI()
        self.notion = NotionClient(notion_api_key)
        
        logger.info("Business Planning Agent initialized successfully")
    
    # ========================================================================
    # MODULE 1: AUTO LEAN CANVAS GENERATOR
    # ========================================================================
    
    async def generate_lean_canvas(
        self,
        idea: str,
        target_market: str = "",
        business_model: str = ""
    ) -> LeanCanvas:
        """
        Generate complete Lean Canvas with all 9 blocks filled automatically
        """
        
        system_prompt = """You are a business strategy expert specializing in Lean Canvas methodology.
Generate a comprehensive Lean Canvas for the given business idea.

The Lean Canvas has 9 blocks:
1. Problem: Top 3 problems to solve
2. Solution: Top 3 features
3. Unique Value Proposition: Single, clear, compelling message
4. Unfair Advantage: Can't be easily copied or bought
5. Customer Segments: Target customers
6. Key Metrics: Key activities to measure
7. Channels: Path to customers
8. Cost Structure: Customer acquisition costs, distribution costs, hosting, people
9. Revenue Streams: Revenue model, lifetime value, revenue, gross margin

Return ONLY valid JSON with this structure:
{
  "problem": {
    "title": "Problem",
    "content": ["problem1", "problem2", "problem3"],
    "description": "Brief description"
  },
  "solution": {
    "title": "Solution",
    "content": ["feature1", "feature2", "feature3"],
    "description": "Brief description"
  },
  "unique_value_proposition": {
    "title": "Unique Value Proposition",
    "content": ["main proposition"],
    "description": "Why you are different and worth buying"
  },
  "unfair_advantage": {
    "title": "Unfair Advantage",
    "content": ["advantage1", "advantage2"],
    "description": "What makes you unique"
  },
  "customer_segments": {
    "title": "Customer Segments",
    "content": ["segment1", "segment2"],
    "description": "Target audience"
  },
  "key_metrics": {
    "title": "Key Metrics",
    "content": ["metric1", "metric2", "metric3"],
    "description": "Key numbers to track"
  },
  "channels": {
    "title": "Channels",
    "content": ["channel1", "channel2", "channel3"],
    "description": "Path to customers"
  },
  "cost_structure": {
    "title": "Cost Structure",
    "content": ["cost1", "cost2", "cost3"],
    "description": "Fixed and variable costs"
  },
  "revenue_streams": {
    "title": "Revenue Streams",
    "content": ["stream1", "stream2"],
    "description": "How you make money"
  }
}"""

        user_prompt = f"""Generate a complete Lean Canvas for this business idea:

IDEA: {idea}
TARGET MARKET: {target_market or "Not specified"}
BUSINESS MODEL: {business_model or "Not specified"}

Fill all 9 blocks of the Lean Canvas with specific, actionable content.
Be concrete and realistic. Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.5,
                json_mode=True,
                max_tokens=3000
            )
            
            data = json.loads(response)
            
            return LeanCanvas(
                problem=LeanCanvasBlock(**data["problem"]),
                solution=LeanCanvasBlock(**data["solution"]),
                unique_value_proposition=LeanCanvasBlock(**data["unique_value_proposition"]),
                unfair_advantage=LeanCanvasBlock(**data["unfair_advantage"]),
                customer_segments=LeanCanvasBlock(**data["customer_segments"]),
                key_metrics=LeanCanvasBlock(**data["key_metrics"]),
                channels=LeanCanvasBlock(**data["channels"]),
                cost_structure=LeanCanvasBlock(**data["cost_structure"]),
                revenue_streams=LeanCanvasBlock(**data["revenue_streams"])
            )
        
        except Exception as e:
            logger.error(f"Error generating Lean Canvas: {str(e)}")
            # Return default canvas
            return self._get_default_lean_canvas()
    
    # ========================================================================
    # MODULE 2: FINANCIAL PROJECTION ESTIMATOR
    # ========================================================================
    
    async def estimate_financials(
        self,
        idea: str,
        business_model: str = "",
        target_market_size: str = ""
    ) -> FinancialEstimate:
        """
        AI estimates 3-year revenue, CAC, and burn rate
        """
        
        system_prompt = """You are a financial analyst and CFO with expertise in startup financial modeling.
Estimate realistic 3-year financial projections for the given business.

Consider:
- Revenue growth trajectory (conservative, realistic estimates)
- Cost structure (fixed and variable costs)
- Customer Acquisition Cost (CAC)
- Lifetime Value (LTV)
- Burn rate and runway
- Break-even analysis

Return ONLY valid JSON:
{
  "projections": [
    {
      "year": 1,
      "revenue": 100000,
      "costs": 150000,
      "profit": -50000,
      "burn_rate": 4167,
      "runway_months": 12
    }
  ],
  "cac": 50,
  "ltv": 500,
  "ltv_cac_ratio": 10,
  "break_even_month": 18,
  "total_funding_needed": 500000,
  "assumptions": ["assumption1", "assumption2"]
}"""

        user_prompt = f"""Estimate 3-year financial projections for this business:

IDEA: {idea}
BUSINESS MODEL: {business_model or "Not specified"}
TARGET MARKET SIZE: {target_market_size or "Not specified"}

Provide:
1. Year-by-year projections (revenue, costs, profit, burn rate, runway)
2. CAC (Customer Acquisition Cost)
3. LTV (Lifetime Value)
4. LTV:CAC ratio
5. Break-even month
6. Total funding needed
7. Key assumptions

Be realistic and conservative. Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.3,
                json_mode=True,
                max_tokens=2000
            )
            
            data = json.loads(response)
            
            projections = []
            for proj in data.get("projections", []):
                projections.append(FinancialProjection(
                    year=proj.get("year", 1),
                    revenue=float(proj.get("revenue", 0)),
                    costs=float(proj.get("costs", 0)),
                    profit=float(proj.get("profit", 0)),
                    burn_rate=float(proj.get("burn_rate", 0)),
                    runway_months=int(proj.get("runway_months", 0))
                ))
            
            return FinancialEstimate(
                projections=projections,
                cac=float(data.get("cac", 0)),
                ltv=float(data.get("ltv", 0)),
                ltv_cac_ratio=float(data.get("ltv_cac_ratio", 0)),
                break_even_month=int(data.get("break_even_month", 0)),
                total_funding_needed=float(data.get("total_funding_needed", 0)),
                assumptions=data.get("assumptions", [])
            )
        
        except Exception as e:
            logger.error(f"Error estimating financials: {str(e)}")
            return self._get_default_financial_estimate()
    
    # ========================================================================
    # MODULE 3: TEAM ROLE MAPPING
    # ========================================================================
    
    async def map_team_roles(
        self,
        idea: str,
        business_model: str = "",
        stage: str = "pre-seed"
    ) -> TeamComposition:
        """
        Suggest team composition based on business model
        """
        
        system_prompt = """You are a startup HR consultant and organizational design expert.
Suggest optimal team composition for the given business at its current stage.

Consider:
- Essential roles for the business model
- Skills required for each role
- Hiring priority and timeline
- Estimated compensation

Return ONLY valid JSON:
{
  "roles": [
    {
      "role": "CTO / Technical Co-Founder",
      "responsibilities": ["resp1", "resp2"],
      "skills_required": ["skill1", "skill2"],
      "priority": "Critical",
      "hiring_timeline": "Month 0"
    }
  ],
  "total_team_size": 5,
  "estimated_payroll_monthly": 50000
}"""

        user_prompt = f"""Suggest team composition for this business:

IDEA: {idea}
BUSINESS MODEL: {business_model or "Not specified"}
STAGE: {stage}

Provide:
1. Key roles needed (5-10 roles)
2. Responsibilities for each role
3. Required skills
4. Hiring priority (Critical/High/Medium/Low)
5. Hiring timeline
6. Total team size
7. Estimated monthly payroll

Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.4,
                json_mode=True,
                max_tokens=2500
            )
            
            data = json.loads(response)
            
            roles = []
            for role_data in data.get("roles", []):
                roles.append(TeamRole(
                    role=role_data.get("role", ""),
                    responsibilities=role_data.get("responsibilities", []),
                    skills_required=role_data.get("skills_required", []),
                    priority=role_data.get("priority", "Medium"),
                    hiring_timeline=role_data.get("hiring_timeline", "")
                ))
            
            return TeamComposition(
                roles=roles,
                total_team_size=int(data.get("total_team_size", 0)),
                estimated_payroll_monthly=float(data.get("estimated_payroll_monthly", 0))
            )
        
        except Exception as e:
            logger.error(f"Error mapping team roles: {str(e)}")
            return self._get_default_team_composition()
    
    # ========================================================================
    # MODULE 4: MARKETING STRATEGY BUILDER
    # ========================================================================
    
    async def build_marketing_strategy(
        self,
        idea: str,
        target_audience: str = "",
        budget: float = 10000
    ) -> MarketingStrategy:
        """
        Generate customer acquisition channels & budgets
        """
        
        system_prompt = """You are a growth marketing expert specializing in customer acquisition.
Design a comprehensive marketing strategy with specific channels, tactics, and budgets.

Return ONLY valid JSON:
{
  "channels": [
    {
      "channel": "Content Marketing",
      "strategy": "Detailed strategy",
      "estimated_cost": 2000,
      "expected_roi": 3.5,
      "timeline": "Months 1-12"
    }
  ],
  "total_budget": 10000,
  "customer_acquisition_strategy": "Overall strategy",
  "retention_strategy": "How to retain customers",
  "growth_tactics": ["tactic1", "tactic2"]
}"""

        user_prompt = f"""Design a marketing strategy for this business:

IDEA: {idea}
TARGET AUDIENCE: {target_audience or "Not specified"}
BUDGET: ${budget}

Provide:
1. Top 5-7 marketing channels
2. Strategy for each channel
3. Estimated cost per channel
4. Expected ROI
5. Timeline
6. Customer acquisition strategy
7. Retention strategy
8. Growth tactics

Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.5,
                json_mode=True,
                max_tokens=2500
            )
            
            data = json.loads(response)
            
            channels = []
            for channel_data in data.get("channels", []):
                channels.append(MarketingChannel(
                    channel=channel_data.get("channel", ""),
                    strategy=channel_data.get("strategy", ""),
                    estimated_cost=float(channel_data.get("estimated_cost", 0)),
                    expected_roi=float(channel_data.get("expected_roi", 0)),
                    timeline=channel_data.get("timeline", "")
                ))
            
            return MarketingStrategy(
                channels=channels,
                total_budget=float(data.get("total_budget", budget)),
                customer_acquisition_strategy=data.get("customer_acquisition_strategy", ""),
                retention_strategy=data.get("retention_strategy", ""),
                growth_tactics=data.get("growth_tactics", [])
            )
        
        except Exception as e:
            logger.error(f"Error building marketing strategy: {str(e)}")
            return self._get_default_marketing_strategy()
    
    # ========================================================================
    # MODULE 5: INVESTOR SUMMARY GENERATOR
    # ========================================================================
    
    async def generate_investor_summary(
        self,
        idea: str,
        lean_canvas: LeanCanvas,
        financials: FinancialEstimate,
        team: TeamComposition
    ) -> InvestorSummary:
        """
        Create a one-page summary with traction & value prop
        """
        
        system_prompt = """You are a venture capital advisor helping founders create compelling investor pitches.
Create a concise, powerful one-page investor summary.

Return ONLY valid JSON:
{
  "elevator_pitch": "30-second pitch",
  "problem_statement": "The problem",
  "solution_overview": "The solution",
  "market_opportunity": "Market size and opportunity",
  "business_model": "How we make money",
  "traction": "Current traction and milestones",
  "competitive_advantage": "Why we'll win",
  "team_highlights": "Team strengths",
  "financial_ask": "How much and what for",
  "use_of_funds": ["use1", "use2"],
  "key_milestones": ["milestone1", "milestone2"]
}"""

        # Prepare context
        context = f"""
IDEA: {idea}

LEAN CANVAS:
- Problem: {', '.join(lean_canvas.problem.content)}
- Solution: {', '.join(lean_canvas.solution.content)}
- UVP: {', '.join(lean_canvas.unique_value_proposition.content)}
- Revenue: {', '.join(lean_canvas.revenue_streams.content)}

FINANCIALS:
- Year 1 Revenue: ${financials.projections[0].revenue if financials.projections else 0}
- CAC: ${financials.cac}
- LTV: ${financials.ltv}
- Funding Needed: ${financials.total_funding_needed}

TEAM:
- Team Size: {team.total_team_size}
- Key Roles: {', '.join([r.role for r in team.roles[:3]])}
"""

        user_prompt = f"""{context}

Create a compelling one-page investor summary that includes:
1. Elevator pitch (30 seconds)
2. Problem statement
3. Solution overview
4. Market opportunity
5. Business model
6. Traction (be realistic for early stage)
7. Competitive advantage
8. Team highlights
9. Financial ask
10. Use of funds
11. Key milestones

Make it compelling and investor-ready. Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.6,
                json_mode=True,
                max_tokens=2000
            )
            
            data = json.loads(response)
            
            return InvestorSummary(
                elevator_pitch=data.get("elevator_pitch", ""),
                problem_statement=data.get("problem_statement", ""),
                solution_overview=data.get("solution_overview", ""),
                market_opportunity=data.get("market_opportunity", ""),
                business_model=data.get("business_model", ""),
                traction=data.get("traction", ""),
                competitive_advantage=data.get("competitive_advantage", ""),
                team_highlights=data.get("team_highlights", ""),
                financial_ask=data.get("financial_ask", ""),
                use_of_funds=data.get("use_of_funds", []),
                key_milestones=data.get("key_milestones", [])
            )
        
        except Exception as e:
            logger.error(f"Error generating investor summary: {str(e)}")
            return self._get_default_investor_summary()
    
    # ========================================================================
    # MODULE 6: REGULATORY COMPLIANCE CHECKER
    # ========================================================================
    
    async def check_regulatory_compliance(
        self,
        idea: str,
        industry: str = "",
        region: str = "United States"
    ) -> RegulatoryCompliance:
        """
        Identify region-specific legal requirements using Groq
        """
        
        system_prompt = """You are a legal compliance expert specializing in startup regulations.
Identify key regulatory requirements for the given business.

Consider:
- Business licenses and permits
- Industry-specific regulations
- Data privacy laws (GDPR, CCPA, etc.)
- Employment laws
- Tax requirements
- Intellectual property
- Consumer protection laws

Return ONLY valid JSON:
{
  "requirements": [
    {
      "category": "Business License",
      "requirement": "Specific requirement",
      "jurisdiction": "State/Federal/Local",
      "priority": "Critical",
      "estimated_cost": 500,
      "timeline": "Before launch",
      "resources": ["resource1", "resource2"]
    }
  ],
  "total_compliance_cost": 5000,
  "critical_deadlines": ["deadline1", "deadline2"],
  "recommended_actions": ["action1", "action2"]
}"""

        user_prompt = f"""Identify regulatory compliance requirements for this business:

IDEA: {idea}
INDUSTRY: {industry or "General"}
REGION: {region}

Provide:
1. Key regulatory requirements (5-10 items)
2. Category (license, data privacy, employment, tax, etc.)
3. Jurisdiction (federal, state, local)
4. Priority level
5. Estimated compliance cost
6. Timeline
7. Resources/links
8. Total compliance cost
9. Critical deadlines
10. Recommended actions

Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.3,
                json_mode=True,
                max_tokens=2500
            )
            
            data = json.loads(response)
            
            requirements = []
            for req_data in data.get("requirements", []):
                requirements.append(ComplianceRequirement(
                    category=req_data.get("category", ""),
                    requirement=req_data.get("requirement", ""),
                    jurisdiction=req_data.get("jurisdiction", ""),
                    priority=req_data.get("priority", "Medium"),
                    estimated_cost=float(req_data.get("estimated_cost", 0)),
                    timeline=req_data.get("timeline", ""),
                    resources=req_data.get("resources", [])
                ))
            
            return RegulatoryCompliance(
                requirements=requirements,
                total_compliance_cost=float(data.get("total_compliance_cost", 0)),
                critical_deadlines=data.get("critical_deadlines", []),
                recommended_actions=data.get("recommended_actions", [])
            )
        
        except Exception as e:
            logger.error(f"Error checking regulatory compliance: {str(e)}")
            return self._get_default_regulatory_compliance()
    
    # ========================================================================
    # MODULE 7: AI CO-FOUNDER CHAT
    # ========================================================================
    
    async def get_cofounder_feedback(
        self,
        business_plan: BusinessPlanResponse,
        focus_area: str = "general"
    ) -> List[CoFounderFeedback]:
        """
        Interactive agent that gives founder-like feedback
        Example: "Your customer segment looks too broad. Narrow it down."
        """
        
        system_prompt = """You are an experienced startup co-founder and advisor.
Review the business plan and provide constructive, actionable feedback.

Focus on:
- Identifying weaknesses or gaps
- Suggesting improvements
- Asking clarifying questions
- Highlighting risks
- Recommending next steps

Return ONLY valid JSON array:
[
  {
    "feedback_type": "suggestion|warning|question",
    "message": "Your feedback message",
    "reasoning": "Why this matters",
    "action_items": ["action1", "action2"]
  }
]"""

        # Prepare business plan summary
        plan_summary = f"""
BUSINESS: {business_plan.business_name}
TAGLINE: {business_plan.tagline}

LEAN CANVAS:
- Problem: {', '.join(business_plan.lean_canvas.problem.content)}
- Solution: {', '.join(business_plan.lean_canvas.solution.content)}
- Customer Segments: {', '.join(business_plan.lean_canvas.customer_segments.content)}

FINANCIALS:
- Year 1 Revenue: ${business_plan.financial_estimate.projections[0].revenue if business_plan.financial_estimate.projections else 0}
- CAC: ${business_plan.financial_estimate.cac}
- LTV:CAC Ratio: {business_plan.financial_estimate.ltv_cac_ratio}

TEAM:
- Team Size: {business_plan.team_composition.total_team_size}
- Monthly Payroll: ${business_plan.team_composition.estimated_payroll_monthly}

MARKETING:
- Budget: ${business_plan.marketing_strategy.total_budget}
- Channels: {', '.join([c.channel for c in business_plan.marketing_strategy.channels[:3]])}
"""

        user_prompt = f"""Review this business plan and provide 3-5 pieces of constructive feedback:

{plan_summary}

FOCUS AREA: {focus_area}

Provide feedback on:
1. Customer segment clarity
2. Financial assumptions
3. Team composition
4. Marketing strategy
5. Overall feasibility

Be direct, honest, and helpful like a co-founder would be. Return as JSON array."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.6,
                json_mode=True,
                max_tokens=2000
            )
            
            data = json.loads(response)
            
            # Handle both array and object with feedback key
            if isinstance(data, dict) and "feedback" in data:
                data = data["feedback"]
            
            feedback_list = []
            for fb_data in data[:5]:
                feedback_list.append(CoFounderFeedback(
                    feedback_type=fb_data.get("feedback_type", "suggestion"),
                    message=fb_data.get("message", ""),
                    reasoning=fb_data.get("reasoning", ""),
                    action_items=fb_data.get("action_items", [])
                ))
            
            return feedback_list
        
        except Exception as e:
            logger.error(f"Error getting co-founder feedback: {str(e)}")
            return []
    
    # ========================================================================
    # MODULE 8: EXPORT TO DOCX
    # ========================================================================
    
    async def export_to_docx(
        self,
        business_plan: BusinessPlanResponse,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export business plan to DOCX format
        """
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping DOCX export")
            return ""
        
        try:
            # Create output path
            if not output_path:
                reports_dir = "reports"
                os.makedirs(reports_dir, exist_ok=True)
                output_path = os.path.join(
                    reports_dir,
                    f"business_plan_{business_plan.plan_id}.docx"
                )
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading(business_plan.business_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tagline
            tagline = doc.add_paragraph(business_plan.tagline)
            tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tagline.runs[0].italic = True
            
            doc.add_paragraph()
            
            # Executive Summary
            doc.add_heading('Executive Summary', 1)
            doc.add_paragraph(business_plan.executive_summary)
            
            # Lean Canvas
            doc.add_heading('Lean Canvas', 1)
            
            canvas_blocks = [
                ('Problem', business_plan.lean_canvas.problem),
                ('Solution', business_plan.lean_canvas.solution),
                ('Unique Value Proposition', business_plan.lean_canvas.unique_value_proposition),
                ('Unfair Advantage', business_plan.lean_canvas.unfair_advantage),
                ('Customer Segments', business_plan.lean_canvas.customer_segments),
                ('Key Metrics', business_plan.lean_canvas.key_metrics),
                ('Channels', business_plan.lean_canvas.channels),
                ('Cost Structure', business_plan.lean_canvas.cost_structure),
                ('Revenue Streams', business_plan.lean_canvas.revenue_streams)
            ]
            
            for block_name, block_data in canvas_blocks:
                doc.add_heading(block_name, 2)
                for item in block_data.content:
                    doc.add_paragraph(item, style='List Bullet')
            
            # Financial Projections
            doc.add_heading('Financial Projections', 1)
            doc.add_paragraph(f"CAC: ${business_plan.financial_estimate.cac}")
            doc.add_paragraph(f"LTV: ${business_plan.financial_estimate.ltv}")
            doc.add_paragraph(f"LTV:CAC Ratio: {business_plan.financial_estimate.ltv_cac_ratio}")
            doc.add_paragraph(f"Break-even Month: {business_plan.financial_estimate.break_even_month}")
            doc.add_paragraph(f"Total Funding Needed: ${business_plan.financial_estimate.total_funding_needed}")
            
            # Team Composition
            doc.add_heading('Team Composition', 1)
            doc.add_paragraph(f"Total Team Size: {business_plan.team_composition.total_team_size}")
            doc.add_paragraph(f"Monthly Payroll: ${business_plan.team_composition.estimated_payroll_monthly}")
            
            for role in business_plan.team_composition.roles:
                doc.add_heading(role.role, 2)
                doc.add_paragraph(f"Priority: {role.priority}")
                doc.add_paragraph(f"Timeline: {role.hiring_timeline}")
            
            # Marketing Strategy
            doc.add_heading('Marketing Strategy', 1)
            doc.add_paragraph(f"Total Budget: ${business_plan.marketing_strategy.total_budget}")
            
            for channel in business_plan.marketing_strategy.channels:
                doc.add_heading(channel.channel, 2)
                doc.add_paragraph(f"Strategy: {channel.strategy}")
                doc.add_paragraph(f"Cost: ${channel.estimated_cost}")
                doc.add_paragraph(f"Expected ROI: {channel.expected_roi}x")
            
            # Investor Summary
            doc.add_heading('Investor Summary', 1)
            doc.add_paragraph(business_plan.investor_summary.elevator_pitch)
            
            # Regulatory Compliance
            doc.add_heading('Regulatory Compliance', 1)
            doc.add_paragraph(f"Total Compliance Cost: ${business_plan.regulatory_compliance.total_compliance_cost}")
            
            for req in business_plan.regulatory_compliance.requirements:
                doc.add_heading(req.category, 2)
                doc.add_paragraph(f"Requirement: {req.requirement}")
                doc.add_paragraph(f"Priority: {req.priority}")
                doc.add_paragraph(f"Cost: ${req.estimated_cost}")
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"DOCX exported: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            return ""
    
    # ========================================================================
    # MODULE 9: EXPORT TO PDF
    # ========================================================================
    
    async def export_to_pdf(
        self,
        business_plan: BusinessPlanResponse,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export business plan to PDF format
        """
        
        if not REPORTLAB_AVAILABLE:
            logger.warning("ReportLab not available, skipping PDF export")
            return ""
        
        try:
            # Create output path
            if not output_path:
                reports_dir = "reports"
                os.makedirs(reports_dir, exist_ok=True)
                output_path = os.path.join(
                    reports_dir,
                    f"business_plan_{business_plan.plan_id}.pdf"
                )
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=28,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            
            tagline_style = ParagraphStyle(
                'Tagline',
                parent=styles['Normal'],
                fontSize=14,
                textColor=colors.HexColor('#666666'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#2563eb'),
                spaceAfter=12,
                spaceBefore=20
            )
            
            # Title Page
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph(business_plan.business_name, title_style))
            story.append(Paragraph(business_plan.tagline, tagline_style))
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(f"Business Plan", styles['Heading3']))
            story.append(Paragraph(f"Generated: {business_plan.created_at}", styles['Normal']))
            story.append(PageBreak())
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", heading_style))
            story.append(Paragraph(business_plan.executive_summary, styles['BodyText']))
            story.append(Spacer(1, 0.3*inch))
            
            # Lean Canvas
            story.append(Paragraph("Lean Canvas", heading_style))
            
            canvas_data = [
                ['Block', 'Content'],
                ['Problem', ', '.join(business_plan.lean_canvas.problem.content[:2])],
                ['Solution', ', '.join(business_plan.lean_canvas.solution.content[:2])],
                ['UVP', ', '.join(business_plan.lean_canvas.unique_value_proposition.content)],
                ['Customer Segments', ', '.join(business_plan.lean_canvas.customer_segments.content[:2])],
                ['Revenue Streams', ', '.join(business_plan.lean_canvas.revenue_streams.content[:2])]
            ]
            
            canvas_table = Table(canvas_data, colWidths=[2*inch, 4*inch])
            canvas_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(canvas_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Financial Projections
            story.append(Paragraph("Financial Projections", heading_style))
            
            fin_data = [
                ['Metric', 'Value'],
                ['CAC', f"${business_plan.financial_estimate.cac}"],
                ['LTV', f"${business_plan.financial_estimate.ltv}"],
                ['LTV:CAC Ratio', f"{business_plan.financial_estimate.ltv_cac_ratio}"],
                ['Break-even Month', f"{business_plan.financial_estimate.break_even_month}"],
                ['Funding Needed', f"${business_plan.financial_estimate.total_funding_needed}"]
            ]
            
            fin_table = Table(fin_data, colWidths=[3*inch, 3*inch])
            fin_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(fin_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Investor Summary
            story.append(Paragraph("Investor Summary", heading_style))
            story.append(Paragraph("<b>Elevator Pitch:</b>", styles['BodyText']))
            story.append(Paragraph(business_plan.investor_summary.elevator_pitch, styles['BodyText']))
            story.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF exported: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            return ""
    
    # ========================================================================
    # MAIN BUSINESS PLANNING PIPELINE
    # ========================================================================
    
    async def create_business_plan(
        self,
        idea: str,
        industry: str = "",
        target_market: str = "",
        business_model: str = "",
        region: str = "United States",
        budget: float = 10000,
        export_formats: List[str] = ["pdf", "docx"]
    ) -> BusinessPlanResponse:
        """
        Main pipeline: Generate complete business plan
        
        Args:
            idea: Business idea description
            industry: Industry/sector
            target_market: Target market description
            business_model: Business model type
            region: Geographic region
            budget: Marketing budget
            export_formats: List of export formats (pdf, docx, notion)
        
        Returns:
            Complete BusinessPlanResponse
        """
        
        plan_id = str(uuid.uuid4())
        logger.info(f"Creating business plan {plan_id} for: {idea[:50]}...")
        
        try:
            # Generate business name and tagline
            business_name, tagline = await self._generate_business_name(idea)
            
            # Run all modules in parallel where possible
            logger.info("Generating Lean Canvas...")
            lean_canvas_task = self.generate_lean_canvas(idea, target_market, business_model)
            
            logger.info("Estimating financials...")
            financials_task = self.estimate_financials(idea, business_model, target_market)
            
            logger.info("Mapping team roles...")
            team_task = self.map_team_roles(idea, business_model)
            
            logger.info("Building marketing strategy...")
            marketing_task = self.build_marketing_strategy(idea, target_market, budget)
            
            logger.info("Checking regulatory compliance...")
            compliance_task = self.check_regulatory_compliance(idea, industry, region)
            
            # Wait for all tasks
            lean_canvas, financials, team, marketing, compliance = await asyncio.gather(
                lean_canvas_task,
                financials_task,
                team_task,
                marketing_task,
                compliance_task
            )
            
            # Generate investor summary (depends on previous results)
            logger.info("Generating investor summary...")
            investor_summary = await self.generate_investor_summary(
                idea, lean_canvas, financials, team
            )
            
            # Generate executive summary
            logger.info("Generating executive summary...")
            executive_summary = await self._generate_executive_summary(
                idea, lean_canvas, financials
            )
            
            # Create business plan response
            business_plan = BusinessPlanResponse(
                plan_id=plan_id,
                business_name=business_name,
                tagline=tagline,
                executive_summary=executive_summary,
                lean_canvas=lean_canvas,
                financial_estimate=financials,
                team_composition=team,
                marketing_strategy=marketing,
                investor_summary=investor_summary,
                regulatory_compliance=compliance,
                co_founder_feedback=[],
                created_at=datetime.now().isoformat()
            )
            
            # Get co-founder feedback
            logger.info("Getting co-founder feedback...")
            business_plan.co_founder_feedback = await self.get_cofounder_feedback(
                business_plan
            )
            
            # Export to requested formats
            if "pdf" in export_formats:
                logger.info("Exporting to PDF...")
                business_plan.pdf_url = await self.export_to_pdf(business_plan)
            
            if "docx" in export_formats:
                logger.info("Exporting to DOCX...")
                business_plan.docx_url = await self.export_to_docx(business_plan)
            
            if "notion" in export_formats:
                logger.info("Exporting to Notion...")
                # Notion export requires parent page ID
                # business_plan.notion_url = await self._export_to_notion(business_plan)
                logger.info("Notion export requires NOTION_API_KEY and parent page ID")
            
            logger.info(f"Business plan {plan_id} created successfully!")
            return business_plan
        
        except Exception as e:
            logger.error(f"Error creating business plan: {str(e)}")
            raise
    
    # ========================================================================
    # HELPER METHODS
    # ========================================================================
    
    async def _generate_business_name(self, idea: str) -> Tuple[str, str]:
        """Generate business name and tagline"""
        
        system_prompt = """You are a branding expert. Generate a catchy business name and tagline.
Return ONLY valid JSON:
{
  "name": "Business Name",
  "tagline": "Catchy tagline"
}"""

        user_prompt = f"""Generate a business name and tagline for this idea:

IDEA: {idea}

Requirements:
- Name should be memorable, unique, and relevant
- Tagline should be concise (5-10 words)
- Both should be professional

Return as JSON."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.8,
                json_mode=True,
                max_tokens=200
            )
            
            data = json.loads(response)
            return data.get("name", "My Startup"), data.get("tagline", "")
        
        except Exception as e:
            logger.error(f"Error generating business name: {str(e)}")
            return "My Startup", "Turning ideas into reality"
    
    async def _generate_executive_summary(
        self,
        idea: str,
        lean_canvas: LeanCanvas,
        financials: FinancialEstimate
    ) -> str:
        """Generate executive summary"""
        
        system_prompt = """You are a business plan writer. Create a compelling executive summary."""

        user_prompt = f"""Create an executive summary for this business:

IDEA: {idea}

LEAN CANVAS:
- Problem: {', '.join(lean_canvas.problem.content)}
- Solution: {', '.join(lean_canvas.solution.content)}
- UVP: {', '.join(lean_canvas.unique_value_proposition.content)}

FINANCIALS:
- Year 1 Revenue: ${financials.projections[0].revenue if financials.projections else 0}
- Funding Needed: ${financials.total_funding_needed}

Write a 2-3 paragraph executive summary that covers:
1. The problem and opportunity
2. The solution and unique value
3. Market potential and financial highlights

Make it compelling and concise."""

        try:
            response = await self.groq.generate(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.6,
                max_tokens=500
            )
            
            return response.strip()
        
        except Exception as e:
            logger.error(f"Error generating executive summary: {str(e)}")
            return "Executive summary unavailable."
    
    def _get_default_lean_canvas(self) -> LeanCanvas:
        """Return default Lean Canvas"""
        return LeanCanvas(
            problem=LeanCanvasBlock("Problem", ["Problem 1", "Problem 2", "Problem 3"]),
            solution=LeanCanvasBlock("Solution", ["Feature 1", "Feature 2", "Feature 3"]),
            unique_value_proposition=LeanCanvasBlock("UVP", ["Unique value proposition"]),
            unfair_advantage=LeanCanvasBlock("Unfair Advantage", ["Advantage 1"]),
            customer_segments=LeanCanvasBlock("Customer Segments", ["Segment 1"]),
            key_metrics=LeanCanvasBlock("Key Metrics", ["Metric 1"]),
            channels=LeanCanvasBlock("Channels", ["Channel 1"]),
            cost_structure=LeanCanvasBlock("Cost Structure", ["Cost 1"]),
            revenue_streams=LeanCanvasBlock("Revenue Streams", ["Revenue 1"])
        )
    
    def _get_default_financial_estimate(self) -> FinancialEstimate:
        """Return default financial estimate"""
        return FinancialEstimate(
            projections=[
                FinancialProjection(1, 0, 0, 0, 0, 0),
                FinancialProjection(2, 0, 0, 0, 0, 0),
                FinancialProjection(3, 0, 0, 0, 0, 0)
            ],
            cac=0,
            ltv=0,
            ltv_cac_ratio=0,
            break_even_month=0,
            total_funding_needed=0,
            assumptions=[]
        )
    
    def _get_default_team_composition(self) -> TeamComposition:
        """Return default team composition"""
        return TeamComposition(
            roles=[],
            total_team_size=0,
            estimated_payroll_monthly=0
        )
    
    def _get_default_marketing_strategy(self) -> MarketingStrategy:
        """Return default marketing strategy"""
        return MarketingStrategy(
            channels=[],
            total_budget=0,
            customer_acquisition_strategy="",
            retention_strategy="",
            growth_tactics=[]
        )
    
    def _get_default_investor_summary(self) -> InvestorSummary:
        """Return default investor summary"""
        return InvestorSummary(
            elevator_pitch="",
            problem_statement="",
            solution_overview="",
            market_opportunity="",
            business_model="",
            traction="",
            competitive_advantage="",
            team_highlights="",
            financial_ask="",
            use_of_funds=[],
            key_milestones=[]
        )
    
    def _get_default_regulatory_compliance(self) -> RegulatoryCompliance:
        """Return default regulatory compliance"""
        return RegulatoryCompliance(
            requirements=[],
            total_compliance_cost=0,
            critical_deadlines=[],
            recommended_actions=[]
        )
    
    def format_response(self, response: BusinessPlanResponse) -> Dict[str, Any]:
        """Format response as dictionary"""
        return asdict(response)


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

async def main():
    """Example usage"""
    
    # Initialize agent
    agent = BusinessPlanningAgent()
    
    # Create business plan
    idea = """
    A mobile app that uses AI to help people learn languages through 
    real-world conversations with native speakers. Users can practice 
    speaking, get instant feedback, and connect with language partners.
    """
    
    business_plan = await agent.create_business_plan(
        idea=idea,
        industry="EdTech",
        target_market="Language learners aged 18-35",
        business_model="Freemium subscription",
        region="United States",
        budget=15000,
        export_formats=["pdf", "docx"]
    )
    
    print(f"\n✓ Business Plan Created: {business_plan.business_name}")
    print(f"  Plan ID: {business_plan.plan_id}")
    print(f"  PDF: {business_plan.pdf_url}")
    print(f"  DOCX: {business_plan.docx_url}")
    print(f"\n{business_plan.executive_summary}")


# ============================================================================
# FASTAPI ROUTER AND ENDPOINTS
# ============================================================================

from fastapi import APIRouter, HTTPException, Depends
from pydantic import BaseModel, Field
from typing import Optional

# Create router
router = APIRouter(prefix="/api/business-plan", tags=["Business Planning"])

# Initialize agent (singleton)
agent = None

def get_agent():
    """Get or create agent instance"""
    global agent
    if agent is None:
        agent = BusinessPlanningAgent()
    return agent


# ============================================================================
# REQUEST/RESPONSE MODELS
# ============================================================================

class BusinessPlanRequest(BaseModel):
    """Business plan creation request"""
    idea: str = Field(..., description="Business idea description", min_length=10)
    industry: Optional[str] = Field("", description="Industry/sector")
    target_market: Optional[str] = Field("", description="Target market description")
    business_model: Optional[str] = Field("", description="Business model type")
    region: Optional[str] = Field("United States", description="Operating region")
    budget: Optional[int] = Field(10000, description="Initial budget")
    export_formats: Optional[List[str]] = Field(["pdf"], description="Export formats")
    
    class Config:
        json_schema_extra = {
            "example": {
                "idea": "AI-powered language learning app with conversation practice",
                "industry": "EdTech",
                "target_market": "Language learners aged 18-35",
                "business_model": "Freemium subscription",
                "region": "United States",
                "budget": 15000,
                "export_formats": ["pdf", "docx"]
            }
        }


# ============================================================================
# API ENDPOINTS
# ============================================================================

@router.post("/create")
async def create_business_plan(request: BusinessPlanRequest):
    """
    Create comprehensive business plan
    
    Generates a complete business plan including:
    - Lean Canvas (all 9 blocks)
    - 3-year financial projections
    - Team composition and hiring timeline
    - Marketing strategy with budget allocation
    - Regulatory compliance checklist
    - Investor summary and funding requirements
    - Exports to PDF/DOCX
    """
    try:
        logger.info(f"Creating business plan for: {request.idea[:50]}...")
        planning_agent = get_agent()
        
        if not planning_agent:
            raise HTTPException(status_code=503, detail="Business Planning Agent not initialized")
        
        # Create business plan
        business_plan: BusinessPlanResponse = await planning_agent.create_business_plan(
            idea=request.idea,
            industry=request.industry or "",
            target_market=request.target_market or "",
            business_model=request.business_model or "",
            region=request.region or "United States",
            budget=request.budget or 10000,
            export_formats=request.export_formats or ["pdf"]
        )
        
        logger.info(f"Business plan created successfully: {business_plan.plan_id}")
        
        # Format response
        formatted_response = planning_agent.format_response(business_plan)
        
        return {
            "status": "success",
            "data": formatted_response
        }
    
    except Exception as e:
        logger.error(f"Error creating business plan: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Business plan creation failed: {str(e)}")


@router.post("/lean-canvas")
async def generate_lean_canvas(request: dict):
    """Generate Lean Canvas only"""
    try:
        logger.info("Generating Lean Canvas...")
        planning_agent = get_agent()
        
        idea = request.get("idea", "")
        target_market = request.get("target_market", "")
        business_model = request.get("business_model", "")
        
        lean_canvas = await planning_agent.generate_lean_canvas(
            idea=idea,
            target_market=target_market,
            business_model=business_model
        )
        
        logger.info("Lean Canvas generated successfully")
        return {
            "status": "success",
            "data": asdict(lean_canvas)
        }
    
    except Exception as e:
        logger.error(f"Error generating lean canvas: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lean Canvas generation failed: {str(e)}")


@router.post("/financials")
async def estimate_financials(request: dict):
    """Estimate financial projections"""
    try:
        logger.info("Estimating financials...")
        planning_agent = get_agent()
        
        idea = request.get("idea", "")
        business_model = request.get("business_model", "")
        target_market_size = request.get("target_market_size", "")
        
        financials = await planning_agent.estimate_financials(
            idea=idea,
            business_model=business_model,
            target_market_size=target_market_size
        )
        
        logger.info("Financials estimated successfully")
        return {
            "status": "success",
            "data": asdict(financials)
        }
    
    except Exception as e:
        logger.error(f"Error estimating financials: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Financial estimation failed: {str(e)}")


@router.post("/team")
async def map_team_roles(request: dict):
    """Map team roles and composition"""
    try:
        logger.info("Mapping team roles...")
        planning_agent = get_agent()
        
        idea = request.get("idea", "")
        business_model = request.get("business_model", "")
        stage = request.get("stage", "pre-seed")
        
        team = await planning_agent.map_team_roles(
            idea=idea,
            business_model=business_model,
            stage=stage
        )
        
        logger.info("Team roles mapped successfully")
        return {
            "status": "success",
            "data": asdict(team)
        }
    
    except Exception as e:
        logger.error(f"Error mapping team roles: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Team mapping failed: {str(e)}")


@router.post("/marketing")
async def build_marketing_strategy(request: dict):
    """Build marketing strategy"""
    try:
        logger.info("Building marketing strategy...")
        planning_agent = get_agent()
        
        idea = request.get("idea", "")
        target_audience = request.get("target_audience", "")
        budget = request.get("budget", 10000)
        
        marketing = await planning_agent.build_marketing_strategy(
            idea=idea,
            target_audience=target_audience,
            budget=budget
        )
        
        logger.info("Marketing strategy built successfully")
        return {
            "status": "success",
            "data": asdict(marketing)
        }
    
    except Exception as e:
        logger.error(f"Error building marketing strategy: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Marketing strategy failed: {str(e)}")


@router.post("/compliance")
async def check_regulatory_compliance(request: dict):
    """Check regulatory compliance"""
    try:
        logger.info("Checking regulatory compliance...")
        planning_agent = get_agent()
        
        idea = request.get("idea", "")
        industry = request.get("industry", "")
        region = request.get("region", "United States")
        
        compliance = await planning_agent.check_regulatory_compliance(
            idea=idea,
            industry=industry,
            region=region
        )
        
        logger.info("Compliance check completed successfully")
        return {
            "status": "success",
            "data": asdict(compliance)
        }
    
    except Exception as e:
        logger.error(f"Error checking compliance: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Compliance check failed: {str(e)}")


@router.get("/health")
async def business_plan_health():
    """Health check for Business Planning Agent"""
    return {
        "status": "ok",
        "agent": "initialized" if agent else "not initialized",
        "timestamp": datetime.now().isoformat()
    }


if __name__ == "__main__":
    asyncio.run(main())
